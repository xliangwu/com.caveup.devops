from docx import Document
from docx.shared import Inches

def process():
    document = Document('测试发现的问题0831.DOCX')

    # tables = document.tables
    # for activeTable in tables:
    #     print(activeTable)
    #
    # document.save('test.docx')

    for p in document.paragraphs:
        print("-----")
        print(p.text)


def process2():
    document = Document('20201102汽车数据月度跟踪.docx')

    tables = document.tables
    for activeTable in tables:
        print(activeTable)
        activeTable._element.getparent().remove(activeTable._element)

    document.save('output2.docx')
    document = Document('output2.docx')

    for p in document.paragraphs:
        r = p.add_run()
        r.add_break()
        r.add_break()
        r.add_picture('t_0.png')

    document.save('output.docx')

if __name__ == '__main__':
    process2()
